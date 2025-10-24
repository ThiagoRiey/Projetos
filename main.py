import os
import platform
import subprocess
from datetime import datetime
from tkinter import Tk, Label, Entry, Text, Button, StringVar, END, messagebox, ttk, Frame
from docx import Document
from openpyxl import Workbook, load_workbook
import matplotlib.pyplot as plt


def abrir_arquivo(caminho):
    sistema = platform.system()
    try:
        if sistema == "Windows":
            os.startfile(caminho)
        elif sistema == "Darwin":
            subprocess.run(["open", caminho])
        else:
            subprocess.run(["xdg-open", caminho])
    except Exception as e:
        messagebox.showwarning("Erro ao abrir arquivo", f"Não foi possível abrir o arquivo:\n{e}")


def append_to_excel(path, header, data):
    if os.path.exists(path):
        wb = load_workbook(path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.append(header)
    ws.append(data)
    wb.save(path)


def replace_placeholders(doc, mapping):
    def replace_text_in_para(para, key, value):
        full_text = "".join(run.text for run in para.runs)
        if f"{{{{{key}}}}}" in full_text:
            new_text = full_text.replace(f"{{{{{key}}}}}", value)
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text

    for p in doc.paragraphs:
        for key, value in mapping.items():
            replace_text_in_para(p, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in mapping.items():
                        replace_text_in_para(p, key, value)
    return doc


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("📋 Gerador e Consulta de Ordens de Serviço")
        self.root.geometry("650x750")
        self.root.resizable(False, False)

       
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("TButton", font=("Segoe UI", 10, "bold"), padding=6)
        style.configure("TEntry", padding=3)

        self.model_path = "Ordem de Serviço.docx"
        self.last_docx = None

        # Variáveis
        self.veiculo = StringVar()
        self.placa = StringVar()
        self.km_entrada = StringVar()
        self.km_saida = StringVar()
        self.step = StringVar()
        self.chave_roda = StringVar()
        self.macaco = StringVar()
        self.triangulo = StringVar()
        self.num_busca = StringVar()

        container = Frame(root, padx=20, pady=15)
        container.pack(fill="both", expand=True)

        ttk.Label(container, text="🚗 Dados do Veículo", font=("Segoe UI", 12, "bold")).grid(row=0, column=0, columnspan=2, pady=(0, 8))

        ttk.Label(container, text="Veículo:").grid(row=1, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.veiculo, width=40).grid(row=1, column=1)

        ttk.Label(container, text="Placa:").grid(row=2, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.placa, width=20).grid(row=2, column=1, sticky="w")

        ttk.Label(container, text="KM na Entrada:").grid(row=3, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.km_entrada, width=20).grid(row=3, column=1, sticky="w")

        ttk.Label(container, text="KM na Saída:").grid(row=4, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.km_saida, width=20).grid(row=4, column=1, sticky="w")

        ttk.Label(container, text="Itens a verificar:").grid(row=5, column=0, sticky="nw", pady=(5, 0))
        self.itens = Text(container, width=45, height=4, font=("Segoe UI", 9))
        self.itens.grid(row=5, column=1, pady=(5, 0))

        ttk.Label(container, text="Observações:").grid(row=6, column=0, sticky="nw", pady=(5, 0))
        self.obs = Text(container, width=45, height=4, font=("Segoe UI", 9))
        self.obs.grid(row=6, column=1, pady=(5, 0))

        ttk.Label(container, text="STEP (S/N):").grid(row=7, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.step, width=5).grid(row=7, column=1, sticky="w")

        ttk.Label(container, text="Chave de Roda (S/N):").grid(row=8, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.chave_roda, width=5).grid(row=8, column=1, sticky="w")

        ttk.Label(container, text="Macaco (S/N):").grid(row=9, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.macaco, width=5).grid(row=9, column=1, sticky="w")

        ttk.Label(container, text="Triângulo (S/N):").grid(row=10, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.triangulo, width=5).grid(row=10, column=1, sticky="w")

        ttk.Button(container, text="💾 Gerar Ordem de Serviço", command=self.generate).grid(row=11, column=0, columnspan=2, pady=10)

        ttk.Separator(container, orient="horizontal").grid(row=12, column=0, columnspan=2, sticky="ew", pady=10)

        ttk.Label(container, text="🔎 Buscar OS nº:", font=("Segoe UI", 10, "bold")).grid(row=13, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.num_busca, width=10).grid(row=13, column=1, sticky="w")
        ttk.Button(container, text="Pesquisar", command=self.buscar_os).grid(row=13, column=1, sticky="e")

        ttk.Button(container, text="📂 Abrir Último Documento", command=self.abrir_ultimo_docx).grid(row=14, column=0, columnspan=2, pady=5)
        ttk.Button(container, text="📊 Gerar Gráfico", command=self.gerar_grafico).grid(row=15, column=0, columnspan=2, pady=5)

        self.label_contador = ttk.Label(container, text=self.get_total_os_text(), font=("Segoe UI", 10, "italic"))
        self.label_contador.grid(row=16, column=0, columnspan=2, pady=(10, 0))

    # === Funções auxiliares ===
    def get_total_os_text(self):
        path = "ordens_servico.xlsx"
        if not os.path.exists(path):
            return "Total de OS: 0"
        wb = load_workbook(path)
        ws = wb.active
        total = ws.max_row - 1
        return f"Total de OS: {max(total, 0)}"

    def next_order_number(self, excel_path):
        if not os.path.exists(excel_path):
            return 1
        wb = load_workbook(excel_path)
        ws = wb.active
        last = ws.max_row
        if last < 2:
            return 1
        num_str = str(ws.cell(row=last, column=1).value)
        if num_str.isdigit():
            return int(num_str) + 1
        return 1

    def generate(self):
        if not os.path.exists(self.model_path):
            messagebox.showerror("Erro", "O arquivo 'Ordem de Serviço.docx' não foi encontrado.")
            return

        os.makedirs("OS_Geradas", exist_ok=True)
        excel_path = "ordens_servico.xlsx"

        numero_os = self.next_order_number(excel_path)
        ano = datetime.now().year
        num_formatado = f"{numero_os:03d}/{ano}"
        data_hoje = datetime.now().strftime("%d/%m/%Y")

        doc = Document(self.model_path)
        mapping = {
            "NUM_OS": num_formatado,
            "DATA_ATUAL": data_hoje,
            "DATA_SOLICITACAO": data_hoje,
            "DATA_DEVOLUCAO": "",
            "VEICULO": self.veiculo.get().strip(),
            "PLACA": self.placa.get().strip(),
            "KM_ENTRADA": self.km_entrada.get().strip(),
            "KM_SAIDA": self.km_saida.get().strip(),
            "ITENS_VERIFICAR": self.itens.get("1.0", END).strip(),
            "OBSERVACOES": self.obs.get("1.0", END).strip(),
            "STEP": self.step.get().upper(),
            "CHAVE_RODA": self.chave_roda.get().upper(),
            "MACACO": self.macaco.get().upper(),
            "TRIANGULO": self.triangulo.get().upper(),
        }

        replace_placeholders(doc, mapping)
        base = f"OS_{num_formatado.replace('/', '-')}_{mapping['PLACA']}"
        docx_path = os.path.join("OS_Geradas", base + ".docx")
        doc.save(docx_path)
        self.last_docx = docx_path

        header = [
            "numero_os", "data", "veiculo", "placa",
            "km_entrada", "km_saida", "itens", "observacoes",
            "step", "chave_roda", "macaco", "triangulo", "docx_path"
        ]
        row = [
            numero_os, data_hoje, mapping["VEICULO"], mapping["PLACA"],
            mapping["KM_ENTRADA"], mapping["KM_SAIDA"], mapping["ITENS_VERIFICAR"],
            mapping["OBSERVACOES"], mapping["STEP"], mapping["CHAVE_RODA"],
            mapping["MACACO"], mapping["TRIANGULO"], docx_path
        ]
        append_to_excel(excel_path, header, row)

        messagebox.showinfo("Sucesso", f"OS nº {num_formatado} criada com sucesso!\n\nArquivo: {docx_path}")
        self.label_contador.config(text=self.get_total_os_text())

        if messagebox.askyesno("Abrir Documento", "Deseja abrir o documento Word agora?"):
            abrir_arquivo(docx_path)

    def buscar_os(self):
        num_str = self.num_busca.get().strip()
        if not num_str.isdigit():
            messagebox.showerror("Erro", "Digite um número válido de OS (somente números).")
            return

        excel_path = "ordens_servico.xlsx"
        if not os.path.exists(excel_path):
            messagebox.showerror("Erro", "Nenhum registro encontrado.")
            return

        wb = load_workbook(excel_path)
        ws = wb.active
        encontrado = None

        for row in ws.iter_rows(min_row=2, values_only=True):
            if str(row[0]) == num_str:
                encontrado = row
                break

        if not encontrado:
            messagebox.showinfo("Resultado", f"Nenhuma OS nº {num_str} encontrada.")
            return

        info = (
            f"Número: {encontrado[0]}\n"
            f"Data: {encontrado[1]}\n"
            f"Veículo: {encontrado[2]}\n"
            f"Placa: {encontrado[3]}\n"
            f"KM Entrada: {encontrado[4]}\n"
            f"KM Saída: {encontrado[5]}\n"
            f"Itens: {encontrado[6]}\n"
            f"Observações: {encontrado[7]}\n"
            f"STEP: {encontrado[8]} | Chave de Roda: {encontrado[9]} | Macaco: {encontrado[10]} | Triângulo: {encontrado[11]}"
        )

        if messagebox.askyesno("OS Encontrada", f"{info}\n\nDeseja abrir o documento Word?"):
            caminho = encontrado[12]
            if caminho and os.path.exists(caminho):
                abrir_arquivo(caminho)
            else:
                messagebox.showwarning("Arquivo não encontrado", "O arquivo Word desta OS não foi localizado.")

    def abrir_ultimo_docx(self):
        if self.last_docx and os.path.exists(self.last_docx):
            abrir_arquivo(self.last_docx)
        else:
            messagebox.showinfo("Aviso", "Nenhum documento recente encontrado.")

    def gerar_grafico(self):
        path = "ordens_servico.xlsx"
        if not os.path.exists(path):
            messagebox.showerror("Erro", "Nenhuma planilha encontrada.")
            return

        wb = load_workbook(path)
        ws = wb.active
        dados = {}

        for row in ws.iter_rows(min_row=2, values_only=True):
            veiculo = row[2]
            if not veiculo:
                continue
            dados[veiculo] = dados.get(veiculo, 0) + 1

        if not dados:
            messagebox.showinfo("Aviso", "Nenhum dado disponível para o gráfico.")
            return

        plt.bar(dados.keys(), dados.values())
        plt.title("Quantidade de OS por Veículo")
        plt.xlabel("Veículo")
        plt.ylabel("Quantidade")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        plt.show()


if __name__ == "__main__":
    root = Tk()
    app = App(root)
    root.mainloop()
